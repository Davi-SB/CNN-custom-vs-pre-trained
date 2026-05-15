"""Gera o documento DOCX resumido do estudo de ablação."""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "documento_ablacao_resumido.docx"
RESULTS_DIR = "results_ablation"


doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

for i in range(1, 4):
    doc.styles[f"Heading {i}"].font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)


def add_p(text, bold=False, italic=False, center=False):
    paragraph = doc.add_paragraph()
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    return paragraph


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            table.rows[row_idx + 1].cells[col_idx].text = str(value)

    doc.add_paragraph()


def add_image(path, width=Inches(5.5)):
    full_path = os.path.join(os.path.dirname(__file__) or ".", path)
    if os.path.exists(full_path):
        doc.add_picture(full_path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_p(f"[Imagem não encontrada: {path}]", italic=True)


title = doc.add_heading("Estudo de Ablação da LightCNN", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_p("Introdução à Aprendizagem Profunda - Estudo de Ablação", bold=True, center=True)
add_p("Componente avaliado: Batch Normalization", center=True)
add_p("PyTorch 2.11.0 + CUDA | NVIDIA GeForce RTX 4050 Laptop GPU", center=True)

doc.add_heading("1. Objetivo", level=1)
add_p(
    "Este estudo avalia o impacto da Batch Normalization na CNN customizada LightCNN. "
    "Foram comparadas duas versões do mesmo modelo: uma com BatchNorm, equivalente à "
    "arquitetura original da APS anterior, e outra sem as camadas BatchNorm2d."
)

doc.add_heading("2. Metodologia", level=1)
add_p(
    "A arquitetura base possui três blocos convolucionais seguidos por um classificador: "
    "Conv2d -> BatchNorm opcional -> ReLU -> MaxPool. O restante da arquitetura foi "
    "mantido igual nas duas variantes."
)

add_table(
    ["Dataset", "Épocas", "Batch size", "Otimizador", "Learning rate"],
    [
        ["MNIST", "5", "128", "Adam", "0,001"],
        ["CIFAR-10", "10", "128", "Adam", "0,001"],
    ],
)

doc.add_heading("3. Resultados", level=1)

doc.add_heading("3.1 MNIST", level=2)
add_table(
    ["Variante", "Parâmetros", "Loss teste", "Acurácia teste", "Tempo"],
    [
        ["Com BatchNorm", "390.858", "0,0242", "99,19%", "241,5s"],
        ["Sem BatchNorm", "390.410", "0,0182", "99,40%", "368,0s"],
    ],
)
add_image(os.path.join(RESULTS_DIR, "mnist_curves.png"))
add_p(
    "No MNIST, a remoção da BatchNorm não prejudicou o desempenho. A versão sem "
    "BatchNorm alcançou acurácia ligeiramente maior, embora tenha levado mais tempo "
    "para treinar."
)

doc.add_heading("3.2 CIFAR-10", level=2)
add_table(
    ["Variante", "Parâmetros", "Loss teste", "Acurácia teste", "Tempo"],
    [
        ["Com BatchNorm", "620.810", "0,7199", "75,11%", "272,8s"],
        ["Sem BatchNorm", "620.362", "0,6740", "76,71%", "181,4s"],
    ],
)
add_image(os.path.join(RESULTS_DIR, "cifar10_curves.png"))
add_p(
    "No CIFAR-10, a versão sem BatchNorm também terminou com acurácia maior no "
    "recorte curto de 10 épocas. A curva com BatchNorm ainda mostra ganho rápido "
    "nas primeiras épocas, indicando possível estabilização inicial."
)

doc.add_heading("3.3 Comparação Final", level=2)
add_image(os.path.join(RESULTS_DIR, "accuracy_bar.png"), width=Inches(5))
add_table(
    ["Dataset", "Melhor variante", "Diferença de acurácia"],
    [
        ["MNIST", "Sem BatchNorm", "+0,21 ponto percentual"],
        ["CIFAR-10", "Sem BatchNorm", "+1,60 ponto percentual"],
    ],
)

doc.add_heading("4. Discussão", level=1)
add_p(
    "O componente escolhido foi a Batch Normalization por estar presente em todos os "
    "blocos convolucionais da LightCNN original. Neste experimento enxuto, sua remoção "
    "não causou degradação de desempenho; pelo contrário, a versão sem BatchNorm foi "
    "ligeiramente melhor nos dois datasets."
)
add_p(
    "Uma explicação possível é que a LightCNN é pequena, usa Adam como otimizador e "
    "mantém Dropout no classificador. Além disso, o estudo foi reduzido em número de "
    "épocas, o que pode favorecer variações de inicialização e ordem dos batches."
)
add_p(
    "Portanto, o resultado não significa que BatchNorm seja inútil em geral, mas sim "
    "que ela não foi essencial para esta arquitetura e configuração específica."
)

doc.add_heading("5. Conclusão", level=1)
add_p(
    "A ablação mostrou que a Batch Normalization não foi indispensável para a LightCNN "
    "neste cenário. A versão sem BatchNorm alcançou 99,40% no MNIST e 76,71% no "
    "CIFAR-10, contra 99,19% e 75,11% da versão com BatchNorm."
)
add_p(
    "Minha opinião é que, para esta CNN pequena e com treinamento curto, a BatchNorm "
    "pode ser removida sem perda de desempenho. Em redes maiores, treinamentos mais "
    "longos ou configurações mais instáveis, ela ainda pode ser útil."
)

doc.save(OUTPUT)
print(f"Documento salvo: {OUTPUT}")
